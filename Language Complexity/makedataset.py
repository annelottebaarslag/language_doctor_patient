import os
import re
import pandas as pd
from docx import Document
from doc2docx import convert

# All lines of abel that should be skipped because they are either an explanation of the file, a comment (between [] or ()) or only existing of numbers (for example 5.10)
def should_skip(text):
    text = text.strip()
    if not text:
        return True
    lower = text.lower()

    prefixes = [
        "consult:",
        "de dikgedrukte",
        "dikgedrukte",
        "(als er",
        "id:",
        "id-",
        "consult",
        "[ opnoemen code",
        "deel ",
        "rechtopstaand",
        "rechtopstaande",
        "schuin",
        "schuine",
        "fragment ",
        "e (schuin)",
        "e ",
        "einde",
        "-einde",

        "broer van",
        "vrouw van",
        "vrouw",
        "relatie tot",
        "zus/vriendin",
        "relatie onbekend",
        "…  van",
        "[",
        "Doctor",
        "tweede Doctor",
        "patiënt",
        "man van",
        "echtgenoot van",
        "de Doctor die",
        "dochter ",
        "man.",
        "vriendin van",
        "supervisor",
        "zoon van",
        "zwager van",
        "partner van",
        "schoonzus van",
        "ex-vrouw",
        "vriend"
    ]

    for pref in prefixes:
        if lower.startswith(pref):
            return True

    # In some sentences, the notation misses a closing or opening round bracket, so we skip those lines
    if text.startswith("(") or text.endswith(")"):
        return True

    # don't include sentences that are only numbers
    if not any(char.isalpha() for char in text):
        return True
    return False

# Make sure the first letter is a capital letter and if no punctuation mark at the end, add a dot.
def format_sentence(text):
    text = text.strip()
    if not text:
        return text
    text = text[0].upper() + text[1:]
    if not re.search(r"[.!?]$", text):
        text += "."
    return text

def remove_brackets(text):
    text = re.sub(r"\([^)]*\)", "", text)
    text = re.sub(r"\[[^\]]*\]", "", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()

# checks if a paragraph is fully italic, except for parts between () 
def paragraph_is_fully_italic(paragraph):
    full_text = "".join(run.text for run in paragraph.runs)
    clean = re.sub(r"\([^)]*\)", "", full_text).strip()
    has_text = bool(clean)

    if not has_text:
        return False

    inside = False
    for run in paragraph.runs:
        extracted = ""
        for char in run.text:
            if char == "(":
                inside = True
                continue
            if char == ")":
                inside = False
                continue
            if not inside:
                extracted += char

        if extracted.strip() and not run.italic:
            return False

    return True

# we want to remove fully bold parts (often comments etc.)
def paragraph_is_fully_bold(paragraph):
    has_text = False
    for run in paragraph.runs:
        t = run.text.strip()
        if not t:
            continue
        has_text = True
        if not run.bold:
            return False
    return has_text

def extract_abel_text(paragraph):
    excluded_pre = ("Vrouw:", "M:", "V:", "Vrouw2:", "E:", "Broer:", "Man:", "A2:", "A:", "V2:", "M2:", "M :", "V :")
    full_text = "".join(run.text for run in paragraph.runs)

    # Delete the parts between () since they are not spoken text
    clean = remove_brackets(full_text).strip()

    # skip if nothing meaningful remains
    if not clean or not any(c.isalpha() for c in clean) or should_skip(clean):
        return None

    if paragraph_is_fully_bold(paragraph):
        return None

    # exclude when in a sentence someone else than the patient is talking
    for pre in excluded_pre:
        idx = clean.find(pre)
        if idx != -1:
            clean = clean[:idx].strip()

    if not clean:
        return None

    is_italic = paragraph_is_fully_italic(paragraph)
    if is_italic:
        label = "Patient"
    else:
        label ="Doctor"

    formatted = format_sentence(clean)
    return formatted, label

def extract_ibis_text(text):
    clean = remove_brackets(text).strip()

    if not clean or should_skip(clean):
        return None

    if clean.startswith(("A:", "A;")):
        content = clean[2:].strip()
        if not content or not any(c.isalpha() for c in content):
            return None
        return format_sentence(content), "Doctor"

    if clean.startswith(("P:", "P;")):
        content = clean[2:].strip()
        if not content or not any(c.isalpha() for c in content):
            return None
        return format_sentence(content), "Patient"

    return None

def process_Abel(doc, dialogue):
    for para in doc.paragraphs:
        res = extract_abel_text(para)
        if res:
            dialogue.append(res)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    res = extract_abel_text(para)
                    if res:
                        dialogue.append(res)


def process_IBIS(doc, dialogue):
    for para in doc.paragraphs:
        for line in para.text.splitlines():
            line = line.strip()
            if not line:
                continue
            res = extract_ibis_text(line)
            if res:
                dialogue.append(res)


def create_dataset(data, goal):
    input_folder = f"Transcripten_Redownload/Transcripten {data}"

    if goal == "classifier":
        output_folder = f"{data}_classifier"
    elif goal == "level":
        output_folder = f"{data}_level"
    elif goal == "word_doc":
        output_folder = f"{data}_word_doc"
    elif goal == "dialogue":
        output_folder = f"{data}_dialogue"
    else:
        return

    os.makedirs(output_folder, exist_ok=True)

    all_arts = []
    all_patients = []

    for entry in sorted(os.listdir(input_folder)):
        entry_path = os.path.join(input_folder, entry)

        # In IBIS, there are some folders inside the folder, and they contain multiple files of the same patient. We take those files together as 1 file.
        if os.path.isdir(entry_path) and data.lower() == "ibis":
            dialogue = []

            for f in os.listdir(entry_path):
                f_path = os.path.join(entry_path, f)
                if f.endswith(".docx"):
                    doc = Document(f_path)
                    process_IBIS(doc, dialogue)

            if goal == "dialogue":
                pd.DataFrame(dialogue, columns=["text", "label"]).to_csv(
                    os.path.join(output_folder, f"{entry}_combined.csv"),
                    index=False,
                    quoting=1
                )

            if goal == "level":
                arts_texts = [t for t, l in dialogue if l == "Doctor"]
                patient_texts = [t for t, l in dialogue if l == "Patient"]
                pd.DataFrame({"text": arts_texts}).to_csv(os.path.join(output_folder, f"{entry}_Arts.csv"), index=False, quoting=1)
                pd.DataFrame({"text": patient_texts}).to_csv(os.path.join(output_folder, f"{entry}_Patient.csv"), index=False, quoting=1)

            all_arts.extend(t for t, l in dialogue if l == "Doctor")
            all_patients.extend(t for t, l in dialogue if l == "Patient")
            continue

        if not entry.endswith(".docx"):
            continue

        dialogue = []
        doc = Document(entry_path)

        if data.lower().startswith("abel"):
            process_Abel(doc, dialogue)
        else:
            process_IBIS(doc, dialogue)

        base = entry.replace(".docx", "")

        if goal == "level":
            arts_texts = [t for t, l in dialogue if l == "Doctor"]
            patient_texts = [t for t, l in dialogue if l == "Patient"]
            pd.DataFrame({"text": arts_texts}).to_csv(
                os.path.join(output_folder, f"{base}_Arts.csv"), index=False, quoting=1)
            pd.DataFrame({"text": patient_texts}).to_csv(
                os.path.join(output_folder, f"{base}_Patient.csv"), index=False, quoting=1)

        if goal == "dialogue":
            pd.DataFrame(dialogue, columns=["text", "label"]).to_csv(
                os.path.join(output_folder, f"{base}.csv"),
                index=False,
                quoting=1
            )

        all_arts.extend(t for t, l in dialogue if l == "Doctor")
        all_patients.extend(t for t, l in dialogue if l == "Patient")

    if goal == "classifier":
        combined_df = pd.DataFrame({"text": all_arts + all_patients,"label": ["Doctor"] * len(all_arts) + ["Patient"] * len(all_patients)})
        combined_df.to_csv(
            os.path.join(output_folder, f"{data}_classifier.csv"),
            index=False,
            quoting=1
        )

    if goal == "word_doc":
        arts_filtered = [t for t in all_arts if 10 <= len(t.split()) <= 30]
        patient_filtered = [t for t in all_patients if 10 <= len(t.split()) <= 30]

        pd.DataFrame({"text": arts_filtered, "label": ["Doctor"]*len(arts_filtered)}).to_csv(
            os.path.join(output_folder, f"{data}_Arts.csv"), index=False, quoting=1
        )
        pd.DataFrame({"text": patient_filtered, "label": ["Patient"]*len(patient_filtered)}).to_csv(
            os.path.join(output_folder, f"{data}_Patient.csv"), index=False, quoting=1
        )


create_dataset("ibis colon", "dialogue")
