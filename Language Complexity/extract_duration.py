# # Run
# create_dataset("ibis")
import os
import re
import pandas as pd
from docx import Document

time_pattern = re.compile(r'\d{1,2}[:.]\d{2}(?::\d{2})?')


def parse_time_to_seconds(t):
    t = t.replace(".", ":")
    parts = t.split(":")

    if len(parts) == 3:
        h, m, s = map(int, parts)
        return h * 3600 + m * 60 + s
    else:
        m, s = map(int, parts)
        return m * 60 + s


def extract_duration(text):

    text = text.replace("\xa0", " ")
    text_clean = text.replace(".", ":").lower()

    # if text ends in -eind
    eind_match = re.search(r'\((\d{1,2}:\d{2})\s*-\s*eind\)', text_clean)
    if eind_match:
        duration_seconds = parse_time_to_seconds(eind_match.group(1))
    else:
        duration_seconds = None

    # text starts with Duur:
    if duration_seconds is None:

        for line in text.split("\n"):
            line = line.strip()

            if not line.lower().startswith("duur:"):
                continue

            line_lower = line.lower()
            line_clean = line.replace(".", ":")
            after_colon = line_clean.split(":", 1)[1].strip()

            if after_colon == "":
                duration_seconds = None
                break

            # time is written in min and sec
            if "min" in line_lower:

                duration_matches = re.findall(r'(\d{1,2}):(\d{2})\s*min', line_clean)

                if duration_matches:
                    duration_seconds = sum(
                        int(m) * 60 + int(s) for m, s in duration_matches
                    )
                    break

                minute_match = re.search(r'(\d+)\s*min', line_lower)
                second_match = re.search(r'(\d+)\s*sec', line_lower)
                hour_match = re.search(r'(\d+)\s*uur', line_lower)

                minutes = int(minute_match.group(1)) if minute_match else 0
                seconds = int(second_match.group(1)) if second_match else 0
                hours = int(hour_match.group(1)) if hour_match else 0

                duration_seconds = hours * 3600 + minutes * 60 + seconds
                break

            # time is written like a sum with outcome (we take the result)
            if "=" in line_clean:
                right_side = line_clean.split("=")[-1].strip()
                times = time_pattern.findall(right_side)

                if times:
                    duration_seconds = parse_time_to_seconds(times[-1])
                    break

            # time is written like a sum (we take the sum)
            if "+" in line_clean:
                total = 0
                parts = line_clean.split("+")

                for part in parts:
                    times = time_pattern.findall(part)
                    for t in times:
                        total += parse_time_to_seconds(t)

                if total > 0:
                    duration_seconds = total
                    break

    
            times = time_pattern.findall(line_clean)

            if times:
                duration_seconds = sum(parse_time_to_seconds(t) for t in times)
                break

            duration_seconds = None
            break

    # no duration
    if duration_seconds is None:

        # time is written at the end
        match = re.search(r'\((\d{1,2}:\d{2})\)\s*$', text_clean.strip())

        if match:
            duration_seconds = parse_time_to_seconds(match.group(1))
        else:
            return "No duration"

    # subtracting of anamnesis etc.
    for line in text.split("\n"):

        line = line.strip().lower()

        if line.startswith("duur:"):
            continue

        line_clean = line.replace(".", ":")

        line_times = time_pattern.findall(line_clean)

        if len(line_times) == 2:
            t1 = parse_time_to_seconds(line_times[0])
            t2 = parse_time_to_seconds(line_times[1])

            diff = abs(t2 - t1)

            if diff > 0:
                duration_seconds -= diff

    return int(round(duration_seconds))


def create_dataset(data):

    duration_list = []
    filename_list = []

    input_folder = f"Transcripten_Redownload/Transcripten {data}"

    for entry in sorted(os.listdir(input_folder)):
        entry_path = os.path.join(input_folder, entry)

        if os.path.isdir(entry_path) and data.lower() == "ibis":

            patient_text = ""

            for f in os.listdir(entry_path):

                if not f.endswith(".docx"):
                    continue

                doc = Document(os.path.join(entry_path, f))

                for para in doc.paragraphs:
                    patient_text += "".join(run.text for run in para.runs) + "\n"

                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                patient_text += "".join(run.text for run in para.runs) + "\n"

            duration = extract_duration(patient_text)

            filename_list.append(entry)
            duration_list.append(duration)
            continue

        if entry.endswith(".docx"):

            doc = Document(entry_path)
            full_text = ""

            for para in doc.paragraphs:
                full_text += "".join(run.text for run in para.runs) + "\n"

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            full_text += "".join(run.text for run in para.runs) + "\n"

            duration = extract_duration(full_text)

            filename_list.append(entry)
            duration_list.append(duration)

    df = pd.DataFrame({
        "filename": filename_list,
        "duration": duration_list
    })

    df.to_csv("durationibiscolon.csv", index=False)

create_dataset("ibis colon")